#!/usr/bin/env python3
"""
build_docx.py — PaglaDESIGN DOCX generator

Converts AGENT_MANUAL.md to DOC/AGENT_MANUAL.docx with proper formatting.
Preserves heading hierarchy, tables, code blocks, and blockquotes.

Usage:
    python bin/build_docx.py
    python bin/build_docx.py --output DOC/AGENT_MANUAL.docx
"""

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor, Cm


# ── Paths ──────────────────────────────────────────────────────────────────

ROOT = Path(__file__).resolve().parent.parent
DEFAULT_INPUT = ROOT / ".ai" / "AGENT_MANUAL.md"
DEFAULT_OUTPUT = ROOT / "DOC" / "AGENT_MANUAL.docx"


# ── Color Palette ──────────────────────────────────────────────────────────

INK = RGBColor(0x0A, 0x0A, 0x0B)       # #0A0A0B
ACCENT = RGBColor(0x6B, 0x7E, 0xFF)     # #6B7EFF
MUTED = RGBColor(0x6B, 0x70, 0x7E)      # #6B707E
CODE_BG = RGBColor(0xF5, 0xF5, 0xF3)    # #F5F5F3
PAGLA_BLUE = RGBColor(0x6B, 0x7E, 0xFF)


# ── Styles ─────────────────────────────────────────────────────────────────

def setup_styles(doc: Document) -> None:
    """Configure document styles for PaglaDESIGN."""

    # -- Normal --
    style = doc.styles["Normal"]
    style.font.name = "Pagla Sans"
    style.font.size = Pt(11)
    style.font.color.rgb = INK
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.5

    # -- Heading 1 --
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Pagla Sans"
    h1.font.size = Pt(28)
    h1.font.bold = True
    h1.font.color.rgb = INK
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.line_spacing = 1.2

    # -- Heading 2 --
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Pagla Sans"
    h2.font.size = Pt(20)
    h2.font.bold = True
    h2.font.color.rgb = PAGLA_BLUE
    h2.paragraph_format.space_before = Pt(20)
    h2.paragraph_format.space_after = Pt(8)
    h2.paragraph_format.line_spacing = 1.2

    # -- Heading 3 --
    h3 = doc.styles["Heading 3"]
    h3.font.name = "Pagla Sans"
    h3.font.size = Pt(14)
    h3.font.bold = True
    h3.font.color.rgb = INK
    h3.paragraph_format.space_before = Pt(16)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.line_spacing = 1.2

    # -- Blockquote --
    if "Blockquote" not in [s.name for s in doc.styles]:
        bq = doc.styles.add_style("Blockquote", WD_STYLE_TYPE.PARAGRAPH)
        bq.font.name = "Pagla Sans"
        bq.font.size = Pt(11)
        bq.font.italic = True
        bq.font.color.rgb = MUTED
        bq.paragraph_format.left_indent = Cm(1.5)
        bq.paragraph_format.space_after = Pt(8)

    # -- Code --
    if "Code" not in [s.name for s in doc.styles]:
        code = doc.styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
        code.font.name = "Consolas"
        code.font.size = Pt(10)
        code.font.color.rgb = INK
        code.paragraph_format.space_before = Pt(4)
        code.paragraph_format.space_after = Pt(4)
        code.paragraph_format.left_indent = Cm(0.5)

    # -- CodeBlock --
    if "CodeBlock" not in [s.name for s in doc.styles]:
        cb = doc.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        cb.font.name = "Consolas"
        cb.font.size = Pt(9)
        cb.font.color.rgb = INK
        cb.paragraph_format.space_before = Pt(4)
        cb.paragraph_format.space_after = Pt(4)
        cb.paragraph_format.left_indent = Cm(1.0)


# ── Parser ─────────────────────────────────────────────────────────────────

def parse_markdown(path: Path) -> list[dict]:
    """Parse markdown into structured blocks."""
    lines = path.read_text(encoding="utf-8").split("\n")
    blocks = []
    i = 0

    while i < len(lines):
        line = lines[i]

        # Blank line
        if line.strip() == "":
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^---+\s*$", line):
            blocks.append({"type": "hr"})
            i += 1
            continue

        # Heading
        m = re.match(r"^(#{1,3})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            blocks.append({"type": f"h{level}", "text": m.group(2)})
            i += 1
            continue

        # Code block
        if line.strip().startswith("```"):
            lang = line.strip().lstrip("`").strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            blocks.append({"type": "code", "lang": lang, "lines": code_lines})
            i += 1
            continue

        # Table
        if "|" in line and i + 1 < len(lines) and re.match(r"^\|[-|:\s]+\|$", lines[i + 1]):
            table_lines = [line]
            i += 2  # skip separator
            while i < len(lines) and "|" in lines[i] and lines[i].strip() != "":
                table_lines.append(lines[i])
                i += 1
            blocks.append({"type": "table", "lines": table_lines})
            continue

        # Blockquote
        if line.startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].startswith(">"):
                quote_lines.append(lines[i].lstrip("> ").strip())
                i += 1
            blocks.append({"type": "blockquote", "text": " ".join(quote_lines)})
            continue

        # List item
        m = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)", line)
        if m:
            indent = len(m.group(1))
            marker = m.group(2)
            text = m.group(3)
            blocks.append({
                "type": "list",
                "indent": indent,
                "marker": marker,
                "text": text,
            })
            i += 1
            continue

        # Paragraph
        para_lines = []
        while i < len(lines) and lines[i].strip() != "" and not lines[i].startswith("#") and not lines[i].startswith(">") and not lines[i].strip().startswith("```") and not re.match(r"^---+\s*$", lines[i]):
            para_lines.append(lines[i])
            i += 1
        if para_lines:
            blocks.append({"type": "paragraph", "text": " ".join(para_lines)})

    return blocks


# ── Inline Formatting ──────────────────────────────────────────────────────

def add_inline_text(paragraph, text: str) -> None:
    """Add text with inline formatting (bold, italic, code, links)."""
    # Pattern: **bold**, *italic*, `code`, [text](url)
    pattern = re.compile(
        r"\*\*(.+?)\*\*"         # bold
        r"|\*(.+?)\*"            # italic
        r"|`([^`]+)`"            # inline code
        r"|\[([^\]]+)\]\([^)]+\)"  # link (text only)
    )

    last_end = 0
    for m in pattern.finditer(text):
        # Add plain text before match
        if m.start() > last_end:
            paragraph.add_run(text[last_end:m.start()])

        if m.group(1):  # bold
            run = paragraph.add_run(m.group(1))
            run.bold = True
        elif m.group(2):  # italic
            run = paragraph.add_run(m.group(2))
            run.italic = True
        elif m.group(3):  # code
            run = paragraph.add_run(m.group(3))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x6B, 0x7E, 0xFF)
        elif m.group(4):  # link
            run = paragraph.add_run(m.group(4))
            run.font.color.rgb = ACCENT

        last_end = m.end()

    # Remaining text
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


# ── Renderer ───────────────────────────────────────────────────────────────

def render_table(doc: Document, lines: list[str]) -> None:
    """Render a markdown table into a DOCX table."""
    rows_data = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        # Skip separator row
        if all(re.match(r"^[-:]+$", c) for c in cells):
            continue
        rows_data.append(cells)

    if not rows_data:
        return

    num_cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.style = "Table Grid"

    # Header row
    for j, cell_text in enumerate(rows_data[0]):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        add_inline_text(p, cell_text)
        p.runs[0].bold = True if p.runs else None
        # Header background
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn("w:shd"), {
            qn("w:fill"): "F5F5F3",
            qn("w:val"): "clear",
        })
        shading.append(shd)

    # Data rows
    for i, row_data in enumerate(rows_data[1:], 1):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.rows[i].cells[j]
                cell.text = ""
                add_inline_text(cell.paragraphs[0], cell_text)


def render_blocks(doc: Document, blocks: list[dict]) -> None:
    """Render parsed blocks into a DOCX document."""
    in_code_block = False
    code_lines = []

    for block in blocks:
        btype = block["type"]

        # Handle code blocks (accumulate, render on hr/end)
        if btype == "code":
            in_code_block = True
            code_lines = block["lines"]
            continue

        if in_code_block:
            # Flush code block
            for cl in code_lines:
                p = doc.add_paragraph(cl, style="CodeBlock")
            in_code_block = False
            code_lines = []

        if btype == "h1":
            p = doc.add_heading(block["text"], level=1)
        elif btype == "h2":
            p = doc.add_heading(block["text"], level=2)
        elif btype == "h3":
            p = doc.add_heading(block["text"], level=3)
        elif btype == "paragraph":
            p = doc.add_paragraph()
            add_inline_text(p, block["text"])
        elif btype == "blockquote":
            p = doc.add_paragraph(block["text"], style="Blockquote")
        elif btype == "list":
            prefix = block["marker"]
            if re.match(r"\d+\.", prefix):
                style = "List Number"
            else:
                style = "List Bullet"
            p = doc.add_paragraph(block["text"], style=style)
            if block["indent"] > 0:
                p.paragraph_format.left_indent = Cm(1.0 + block["indent"] * 0.5)
        elif btype == "table":
            render_table(doc, block["lines"])
        elif btype == "hr":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            # Thin line via border
            pPr = p._element.get_or_add_pPr()
            pBdr = pPr.makeelement(qn("w:pBdr"), {})
            bottom = pBdr.makeelement(qn("w:bottom"), {
                qn("w:val"): "single",
                qn("w:sz"): "4",
                qn("w:space"): "1",
                qn("w:color"): "E8E8E6",
            })
            pBdr.append(bottom)
            pPr.append(pBdr)

    # Flush trailing code block
    if in_code_block:
        for cl in code_lines:
            p = doc.add_paragraph(cl, style="CodeBlock")


# ── Footer ─────────────────────────────────────────────────────────────────

def add_footer(doc: Document) -> None:
    """Add footer with document info."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PaglaDESIGN · Agent Manual · ")
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED

    run2 = p.add_run("Intelligence, Unhinged.")
    run2.font.size = Pt(8)
    run2.font.italic = True
    run2.font.color.rgb = ACCENT


# ── Main ───────────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(description="Generate AGENT_MANUAL.docx")
    parser.add_argument(
        "--input", "-i",
        type=Path,
        default=DEFAULT_INPUT,
        help="Input markdown file",
    )
    parser.add_argument(
        "--output", "-o",
        type=Path,
        default=DEFAULT_OUTPUT,
        help="Output docx file",
    )
    args = parser.parse_args()

    # Ensure output directory exists
    args.output.parent.mkdir(parents=True, exist_ok=True)

    # Parse
    blocks = parse_markdown(args.input)

    # Build document
    doc = Document()
    setup_styles(doc)

    # Set margins
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    render_blocks(doc, blocks)
    add_footer(doc)

    # Save
    doc.save(str(args.output))
    print(f"Generated: {args.output}")


if __name__ == "__main__":
    main()
