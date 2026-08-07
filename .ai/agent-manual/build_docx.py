#!/usr/bin/env python3
"""
PaglaDESIGN Agent Manual — DOCX Build Script
==============================================
Compiles AGENT_MANUAL.md into a Word document styled with the canonical
PaglaDESIGN monochrome tokens.

Usage:
    python build_docx.py                    # Generate to DOC/AGENT_MANUAL.docx
    python build_docx.py --output PATH      # Generate to custom path
    python build_docx.py --verify           # Verify markdown structure only

Dependencies:
    pip install python-docx

Output:
    DOC/AGENT_MANUAL.docx (or custom path via --output)
"""

import re
import sys
import os
import argparse
from pathlib import Path
from datetime import datetime

# Force UTF-8 output (ecosystem convention; keep box-drawing/emoji intact)
sys.stdout.reconfigure(encoding="utf-8", errors="replace")
sys.stderr.reconfigure(encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt, Inches, Cm, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# =============================================================================
# CONFIGURATION
# =============================================================================

BASE_DIR = Path(__file__).resolve().parent
MARKDOWN_SOURCE = BASE_DIR / "AGENT_MANUAL.md"
DEFAULT_OUTPUT = BASE_DIR / "DOC" / "AGENT_MANUAL.docx"

# Document metadata
META = {
    "title": "PaglaDESIGN Agent Operating Manual",
    "subtitle": "Agent Operating Manual",
    "version": "2.0.0",
    "date": "2026-08-07",
    "license": "© 2026 AYNAGHOR. Intelligence, Unhinged.",
}

# Color palette — canonical PaglaDESIGN tokens (monochrome, accent meaning-only)
COLORS = {
    "primary": "0A0A0B",      # color.base.ink — headings and body
    "body": "0A0A0B",         # color.base.ink
    "secondary": "6B707E",    # color.base.muted
    "accent": "6B7EFF",       # color.accent.primary — meaning only
    "surface": "F5F5F3",      # color.base.surface — table banding
    "border": "E8E8E6",       # color.base.border
    "white": "FFFFFF",        # color.base.paper
    "black": "000000",
}

# Font configuration (document fonts; web surfaces use PaglaAI Sans via base.css)
FONTS = {
    "heading_ascii": "Calibri",
    "heading_eastAsia": "Microsoft YaHei",
    "body_ascii": "Calibri",
    "body_eastAsia": "Microsoft YaHei",
    "mono": "Consolas",
}


# =============================================================================
# MARKDOWN PARSER
# =============================================================================

class MarkdownNode:
    """Base class for parsed markdown nodes."""
    pass


class HeadingNode(MarkdownNode):
    def __init__(self, level, text, anchor=None):
        self.level = level  # 1-6
        self.text = text
        self.anchor = anchor or re.sub(r'[^\w-]', '', text.lower().replace(" ", "-"))


class ParagraphNode(MarkdownNode):
    def __init__(self, text, bold=False, italic=False):
        self.text = text
        self.bold = bold
        self.italic = italic


class TableNode(MarkdownNode):
    def __init__(self, headers, rows, caption=None):
        self.headers = headers
        self.rows = rows
        self.caption = caption


class CodeBlockNode(MarkdownNode):
    def __init__(self, code, language=""):
        self.code = code
        self.language = language


class BlockquoteNode(MarkdownNode):
    def __init__(self, text):
        self.text = text


class RuleNode(MarkdownNode):
    """Horizontal rule."""
    pass


class ListNode(MarkdownNode):
    def __init__(self, items, ordered=False, level=0):
        self.items = items
        self.ordered = ordered
        self.level = level


def parse_markdown(filepath: str) -> list:
    """
    Parse markdown file into a list of MarkdownNode objects.
    
    Handles:
    - Headings (# ## ### ####)
    - Paragraphs
    - Tables (pipe-delimited)
    - Code blocks (fenced with ```)
    - Blockquotes (>)
    - Horizontal rules (---)
    - Lists (- or * or numbered)
    - Bold/italic inline formatting
    """
    with open(filepath, "r", encoding="utf-8") as f:
        lines = f.readlines()
    
    nodes = []
    i = 0
    in_code_block = False
    code_lines = []
    code_language = ""
    in_table = False
    table_rows = []
    table_headers = []
    
    while i < len(lines):
        line = lines[i].rstrip("\n")
        
        # Code block handling
        if line.startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_language = line[3:].strip()
                code_lines = []
            else:
                in_code_block = False
                nodes.append(CodeBlockNode("\n".join(code_lines), code_language))
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Skip empty lines (but use as paragraph breaks)
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        
        # Headings
        if stripped.startswith("#"):
            level = 0
            for ch in stripped:
                if ch == "#":
                    level += 1
                else:
                    break
            text = stripped[level:].strip()
            # Remove anchor links like [text](#anchor)
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            nodes.append(HeadingNode(min(level, 6), text))
            i += 1
            continue
        
        # Horizontal rule
        if re.match(r'^-{3,}$', stripped) or re.match(r'^\*{3,}$', stripped):
            nodes.append(RuleNode())
            i += 1
            continue
        
        # Table detection (starts with |)
        if stripped.startswith("|") and not stripped.startswith("|--"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            
            # Check if next line is separator
            if i + 1 < len(lines) and re.match(r'^[\| :\-]+$', lines[i+1].strip()):
                table_headers = cells
                i += 2  # Skip header separator
                table_rows = []
                
                # Read table rows
                while i < len(lines):
                    row_line = lines[i].rstrip("\n").strip()
                    if row_line.startswith("|"):
                        row_cells = [c.strip() for c in row_line.strip("|").split("|")]
                        # Skip separator rows
                        if not re.match(r'^[\| :\-]+$', row_line):
                            table_rows.append(row_cells)
                        i += 1
                    else:
                        break
                
                nodes.append(TableNode(table_headers, table_rows))
                in_table = False
                continue
            else:
                # Inline table-like line, treat as paragraph
                nodes.append(ParagraphNode(stripped))
                i += 1
                continue
        
        # Blockquote
        if stripped.startswith(">"):
            quote_text = stripped.lstrip("> ").strip()
            nodes.append(BlockquoteNode(quote_text))
            i += 1
            continue
        
        # List items
        if stripped.startswith("- ") or stripped.startswith("* "):
            items = []
            while i < len(lines) and (lines[i].strip().startswith("- ") or lines[i].strip().startswith("* ")):
                item_text = lines[i].strip()[2:].strip()
                # Handle inline bold/italic
                item_text = process_inline_formatting(item_text)
                items.append(item_text)
                i += 1
            nodes.append(ListNode(items, ordered=False))
            continue
        
        # Numbered list
        list_match = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if list_match:
            items = []
            while i < len(lines) and re.match(r'^\d+\.\s+', lines[i].strip()):
                item_text = re.sub(r'^\d+\.\s+', '', lines[i].strip())
                item_text = process_inline_formatting(item_text)
                items.append(item_text)
                i += 1
            nodes.append(ListNode(items, ordered=True))
            continue
        
        # Regular paragraph — collect consecutive non-empty lines
        para_lines = []
        while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith("#") and not lines[i].strip().startswith("|") and not lines[i].strip().startswith("```") and not lines[i].strip().startswith(">") and not lines[i].strip().startswith("- ") and not re.match(r'^[\-\*]{3,}$', lines[i].strip()) and not re.match(r'^\d+\.\s+', lines[i].strip()):
            para_lines.append(lines[i].rstrip("\n"))
            i += 1
        
        if para_lines:
            para_text = " ".join(para_lines).strip()
            para_text = process_inline_formatting(para_text)
            nodes.append(ParagraphNode(para_text))
    
    return nodes


def process_inline_formatting(text: str) -> str:
    """
    Process inline markdown formatting.
    Returns cleaned text (formatting markers removed for now;
    docx builder handles bold/italic via runs).
    """
    # Remove bold markers but note them
    # Remove italic markers
    # Remove inline code markers
    text = text.replace("**", "").replace("__", "")
    text = text.replace("*", "_").replace("_", "")
    text = text.replace("`", "")
    # Remove links but keep text
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    return text


# =============================================================================
# DOCUMENT BUILDER
# =============================================================================

class DocxBuilder:
    """Builds a Word document from parsed markdown nodes."""
    
    def __init__(self, output_path: Path):
        self.output_path = output_path
        self.doc = Document()
        self._setup_styles()
        self._setup_page()
        
        # Statistics tracking
        self.stats = {
            "paragraphs": 0,
            "tables": 0,
            "headings_1": 0,
            "headings_2": 0,
            "headings_3": 0,
        }
    
    def _setup_styles(self):
        """Configure document styles per PaglaDESIGN design system."""
        styles = self.doc.styles
        
        # Normal style (body text)
        normal = styles["Normal"]
        normal.font.name = FONTS["body_ascii"]
        normal.font.size = Pt(11)
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONTS["body_eastAsia"])
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        normal.paragraph_format.space_after = Pt(8)
        normal.paragraph_format.first_line_indent = Cm(0.75)
        
        # Heading 1 (H2 in markdown → main sections)
        h1 = styles["Heading 1"]
        h1.font.name = FONTS["heading_ascii"]
        h1.font.size = Pt(18)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor.from_string(COLORS["primary"])
        h1._element.rPr.rFonts.set(qn("w:eastAsia"), FONTS["heading_eastAsia"])
        h1.paragraph_format.space_before = Pt(24)
        h1.paragraph_format.space_after = Pt(12)
        h1.paragraph_format.first_line_indent = Cm(0)
        
        # Heading 2 (H3 in markdown → subsections)
        h2 = styles["Heading 2"]
        h2.font.name = FONTS["heading_ascii"]
        h2.font.size = Pt(14)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor.from_string(COLORS["primary"])
        h2._element.rPr.rFonts.set(qn("w:eastAsia"), FONTS["heading_eastAsia"])
        h2.paragraph_format.space_before = Pt(18)
        h2.paragraph_format.space_after = Pt(8)
        h2.paragraph_format.first_line_indent = Cm(0)
        
        # Heading 3 (H4 in markdown → sub-subsections)
        h3 = styles["Heading 3"]
        h3.font.name = FONTS["heading_ascii"]
        h3.font.size = Pt(12)
        h3.font.bold = True
        h3.font.color.rgb = RGBColor.from_string(COLORS["secondary"])
        h3._element.rPr.rFonts.set(qn("w:eastAsia"), FONTS["heading_eastAsia"])
        h3.paragraph_format.space_before = Pt(12)
        h3.paragraph_format.space_after = Pt(6)
        h3.paragraph_format.first_line_indent = Cm(0)
        
        # Code style (for code blocks)
        try:
            code_style = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        except:
            code_style = styles.get("CodeBlock", styles["Normal"])
        code_style.font.name = FONTS["mono"]
        code_style.font.size = Pt(9)
        code_style.font.color.rgb = RGBColor.from_string(COLORS["body"])
        code_style.paragraph_format.left_indent = Cm(1)
        code_style.paragraph_format.right_indent = Cm(0.5)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
        code_style.paragraph_format.first_line_indent = Cm(0)
        
        # Quote style
        try:
            quote_style = styles.add_style("BlockQuote", WD_STYLE_TYPE.PARAGRAPH)
        except:
            quote_style = styles.get("BlockQuote", styles["Normal"])
        quote_style.font.name = FONTS["body_ascii"]
        quote_style.font.size = Pt(10)
        quote_style.font.italic = True
        quote_style.font.color.rgb = RGBColor.from_string(COLORS["secondary"])
        quote_style._element.rPr.rFonts.set(qn("w:eastAsia"), FONTS["body_eastAsia"])
        quote_style.paragraph_format.left_indent = Cm(1.5)
        quote_style.paragraph_format.right_indent = Cm(1)
        quote_style.paragraph_format.space_before = Pt(8)
        quote_style.paragraph_format.space_after = Pt(8)
        quote_style.paragraph_format.first_line_indent = Cm(0)
    
    def _setup_page(self):
        """Configure page layout."""
        section = self.doc.sections[0]
        section.page_width = Inches(8.5)   # A4-ish
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1)
    
    def add_cover_page(self):
        """Add a professional cover page."""
        # Add some spacing at top
        for _ in range(4):
            self.doc.add_paragraph()
        
        # Title
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(META["title"])
        run.font.name = FONTS["heading_ascii"]
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(COLORS["primary"])
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONTS["heading_eastAsia"])
        
        # Subtitle
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(META["subtitle"])
        run.font.name = FONTS["heading_ascii"]
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor.from_string(COLORS["secondary"])
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONTS["heading_eastAsia"])
        
        # Version badge
        version_p = self.doc.add_paragraph()
        version_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = version_p.add_run(f"Version {META['version']}")
        run.font.name = FONTS["heading_ascii"]
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor.from_string(COLORS["accent"])
        
        # Spacing
        for _ in range(6):
            self.doc.add_paragraph()
        
        # Metadata table
        meta_data = [
            ("Status", "Adopted (Decision D-020)"),
            ("Date", META["date"]),
            ("License", META["license"]),
        ]
        
        meta_table = self.doc.add_table(rows=len(meta_data), cols=2)
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for idx, (label, value) in enumerate(meta_data):
            row = meta_table.rows[idx]
            
            label_cell = row.cells[0]
            label_cell.text = label
            label_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in label_cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor.from_string(COLORS["secondary"])
            
            value_cell = row.cells[1]
            value_cell.text = value
            value_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in value_cell.paragraphs[0].runs:
                run.font.size = Pt(10)
        
        # Page break after cover
        self.doc.add_page_break()
    
    def build(self, nodes: list):
        """Build document from parsed nodes."""
        self.add_cover_page()
        
        # Add Table of Contents placeholder
        toc_heading = self.doc.add_heading("Table of Contents", level=1)
        toc_para = self.doc.add_paragraph()
        toc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = toc_para.add_run("(Right-click and select 'Update Field' to populate)")
        run.font.italic = True
        run.font.color.rgb = RGBColor.from_string(COLORS["secondary"])
        run.font.size = Pt(9)
        
        # Add TOC field
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        fldChar1 = parse_xml(r'<w:fldChar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fldCharType="begin"/>')
        run._r.append(fldChar1)
        
        run = paragraph.add_run()
        instrText = parse_xml(r'<w:instrText xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> TOC \o "1-3" \h \z \u </w:instrText>')
        run._r.append(instrText)
        
        run = paragraph.add_run()
        fldChar2 = parse_xml(r'<w:fldChar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fldCharType="separate"/>')
        run._r.append(fldChar2)
        
        run = paragraph.add_run()
        fldChar3 = parse_xml(r'<w:fldChar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fldCharType="end"/>')
        run._r.append(fldChar3)
        
        self.doc.add_page_break()
        
        # Process all nodes
        for node in nodes:
            self._process_node(node)
        
        return self
    
    def _process_node(self, node: MarkdownNode):
        """Dispatch node to appropriate handler."""
        if isinstance(node, HeadingNode):
            self._add_heading(node)
        elif isinstance(node, ParagraphNode):
            self._add_paragraph(node)
        elif isinstance(node, TableNode):
            self._add_table(node)
        elif isinstance(node, CodeBlockNode):
            self._add_code_block(node)
        elif isinstance(node, BlockquoteNode):
            self._add_blockquote(node)
        elif isinstance(node, ListNode):
            self._add_list(node)
        elif isinstance(node, RuleNode):
            self._add_horizontal_rule()
    
    def _add_heading(self, node: HeadingNode):
        """Add a heading."""
        # Map markdown H1 (document title) → skip (cover has it)
        # markdown H2 (##) → Word Heading 1
        # markdown H3 (###) → Word Heading 2
        # markdown H4 (####) → Word Heading 3
        
        if node.level == 1:
            return  # Document title already on cover
        
        word_level = min(node.level - 1, 3)  # Map H2→1, H3→2, H4→3
        
        heading = self.doc.add_heading(node.text, level=word_level)
        
        # Track statistics
        if word_level == 1:
            self.stats["headings_1"] += 1
        elif word_level == 2:
            self.stats["headings_2"] += 1
        else:
            self.stats["headings_3"] += 1
    
    def _add_paragraph(self, node: ParagraphNode):
        """Add a paragraph."""
        para = self.doc.add_paragraph()
        
        # Check if starts with common list-like patterns that should be plain text
        text = node.text
        
        run = para.add_run(text)
        run.font.name = FONTS["body_ascii"]
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor.from_string(COLORS["body"])
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONTS["body_eastAsia"])
        
        # Apply alignment based on content
        if text.startswith("```") or text.startswith("┌"):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self.stats["paragraphs"] += 1
    
    def _add_table(self, node: TableNode):
        """Add a formatted table."""
        # Add caption if present
        if node.caption:
            cap = self.doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = cap.add_run(node.caption)
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor.from_string(COLORS["secondary"])
        
        # Calculate column count
        num_cols = max(len(node.headers), max((len(row) for row in node.rows), default=0))
        if num_cols == 0:
            return
        
        # Create table
        table = self.doc.add_table(rows=1 + len(node.rows), cols=num_cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        header_row = table.rows[0]
        for idx, header_text in enumerate(node.headers):
            if idx < num_cols:
                cell = header_row.cells[idx]
                cell.text = header_text
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Style header
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor.from_string(COLORS["white"])
                # Header shading
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{COLORS["primary"]}" w:val="clear"/>')
                cell._tc.get_or_add_tcPr().append(shading)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        # Data rows
        for row_idx, row_data in enumerate(node.rows):
            row = table.rows[row_idx + 1]
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < num_cols:
                    cell = row.cells[col_idx]
                    cell.text = cell_text
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in cell.paragraphs[0].runs:
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor.from_string(COLORS["body"])
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    
                    # Alternating row colors
                    if row_idx % 2 == 1:
                        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{COLORS["surface"]}" w:val="clear"/>')
                        cell._tc.get_or_add_tcPr().append(shading)
        
        # Set column widths to auto
        for row in table.rows:
            for cell in row.cells:
                cell.width = Inches(1.5)
        
        self.stats["tables"] += 1
    
    def _add_code_block(self, node: CodeBlockNode):
        """Add a code block."""
        para = self.doc.add_paragraph()
        para.style = "CodeBlock"
        
        run = para.add_run(node.code)
        run.font.name = FONTS["mono"]
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(COLORS["body"])
        
        # Add light background shading via paragraph border simulation
        # (python-docx limitation: no direct paragraph background)
    
    def _add_blockquote(self, node: BlockquoteNode):
        """Add a blockquote."""
        para = self.doc.add_paragraph()
        para.style = "BlockQuote"
        
        run = para.add_run(node.text)
        run.font.italic = True
        run.font.color.rgb = RGBColor.from_string(COLORS["secondary"])
    
    def _add_list(self, node: ListNode):
        """Add a bulleted or numbered list."""
        for item_text in node.items:
            if node.ordered:
                para = self.doc.add_paragraph(item_text, style="List Number")
            else:
                para = self.doc.add_paragraph(item_text, style="List Bullet")
            
            for run in para.runs:
                run.font.name = FONTS["body_ascii"]
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor.from_string(COLORS["body"])
                run._element.rPr.rFonts.set(qn("w:eastAsia"), FONTS["body_eastAsia"])
    
    def _add_horizontal_rule(self):
        """Add a horizontal rule (page break for major sections)."""
        # Add a thin-line paragraph border simulation
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("─" * 50)
        run.font.color.rgb = RGBColor.from_string(COLORS["secondary"])
        run.font.size = Pt(8)
    
    def save(self, path: Path = None):
        """Save document to file."""
        save_path = path or self.output_path
        save_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(save_path))
        return save_path
    
    def get_stats(self) -> dict:
        """Return build statistics."""
        return self.stats.copy()


# =============================================================================
# MAIN EXECUTION
# =============================================================================

def verify_source(path: Path) -> bool:
    """Verify markdown source exists and has expected structure."""
    if not path.exists():
        print(f"❌ Source file not found: {path}")
        return False
    
    with open(path, "r", encoding="utf-8") as f:
        content = f.read()
    
    # Count headings
    h2_count = len(re.findall(r"^## ", content, re.MULTILINE))
    h3_count = len(re.findall(r"^### ", content, re.MULTILINE))
    
    print(f"✅ Source verified: {path}")
    print(f"   Size: {len(content):,} characters")
    print(f"   H2 sections: {h2_count}")
    print(f"   H3 subsections: {h3_count}")
    
    # Expected sections check
    expected_sections = [
        "Overview",
        "Working Relationship",
        "Token Inheritance",
        "Component Inheritance",
        "Layout & Templates",
        "Prompt Formula",
        "STITCH Integration",
        "Lint Discipline",
        "Skills",
        "MCP Tools",
        "Build Playbook",
        "Anti-Patterns",
    ]
    
    found = sum(1 for s in expected_sections if s.lower() in content.lower())
    print(f"   Expected sections found: {found}/{len(expected_sections)}")
    
    return found == len(expected_sections)


def main():
    parser = argparse.ArgumentParser(
        description="Build PaglaDESIGN Agent Manual DOCX from markdown source"
    )
    parser.add_argument(
        "--output", "-o",
        type=Path,
        default=DEFAULT_OUTPUT,
        help=f"Output DOCX path (default: {DEFAULT_OUTPUT})"
    )
    parser.add_argument(
        "--verify", "-v",
        action="store_true",
        help="Only verify source structure, don't build"
    )
    parser.add_argument(
        "--source", "-s",
        type=Path,
        default=MARKDOWN_SOURCE,
        help=f"Source markdown path (default: {MARKDOWN_SOURCE})"
    )
    
    args = parser.parse_args()
    
    print("=" * 60)
    print("PaglaDESIGN Agent Manual — DOCX Builder")
    print("=" * 60)
    print()
    
    # Verify source
    if not verify_source(args.source):
        sys.exit(1)
    
    if args.verify:
        print("\n✅ Verification complete. Use without --verify to build.")
        sys.exit(0)
    
    # Parse markdown
    print(f"\n📖 Parsing markdown: {args.source}")
    nodes = parse_markdown(str(args.source))
    print(f"   Parsed {len(nodes)} nodes")
    
    # Build document
    print(f"\n🔨 Building document...")
    builder = DocxBuilder(args.output)
    builder.build(nodes)
    
    # Save
    output_path = builder.save(args.output)
    
    # Report stats
    stats = builder.get_stats()
    print(f"\n✅ Document saved: {output_path}")
    print(f"   File size: {output_path.stat().st_size / 1024:.1f} KB")
    print(f"\n📊 Document Statistics:")
    print(f"   Paragraphs:     {stats['paragraphs']}")
    print(f"   Tables:         {stats['tables']}")
    print(f"   Heading 1 (§):  {stats['headings_1']}")
    print(f"   Heading 2 (§§): {stats['headings_2']}")
    print(f"   Heading 3 (§§§):{stats['headings_3']}")
    print()
    print("=" * 60)
    print("Build complete! ✨")
    print("=" * 60)


if __name__ == "__main__":
    main()
